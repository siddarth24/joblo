import os
import re
from PIL import Image
import pytesseract
import pdfplumber
from docx import Document
from pdf2image import convert_from_path
import docx2txt

def extract_text_and_links_from_file(file_path):
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    file_extension = os.path.splitext(file_path)[1].lower()
    text = ""
    links = []

    # Extract text and links based on file extension
    if file_extension == ".pdf":
        # First attempt to extract text normally
        text, links = extract_text_and_links_from_pdf(file_path)
        # Fallback to OCR if no text was extracted
        if not text.strip():
            ocr_text, ocr_links = extract_text_and_links_from_pdf_ocr(file_path)
            text += ocr_text
            links += ocr_links
    elif file_extension == ".docx":
        text, links = extract_text_and_links_from_docx(file_path)
    elif file_extension == ".txt":
        text, links = extract_text_and_links_from_txt(file_path)
    elif file_extension in [".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".gif"]:
        text, links = extract_text_and_links_from_image(file_path)
    else:
        raise ValueError(f"Unsupported file format: {file_extension}")
    
    # Clean the extracted text
    cleaned_text = clean_text(text)
    return cleaned_text, links

def extract_text_and_links_from_pdf_ocr(pdf_path):
    """Extracts text from PDF using OCR for image-based PDFs."""
    text = ""
    links = []
    try:
        # Convert each PDF page to an image
        images = convert_from_path(pdf_path)
        for image in images:
            # Extract text from the image using pytesseract
            page_text = pytesseract.image_to_string(image)
            text += page_text + "\n"
            # Extract links from the OCR text
            page_links = re.findall(r'https?://\S+', page_text)
            links.extend(page_links)
    except Exception as e:
        raise RuntimeError(f"Error during PDF OCR extraction: {e}")
    return text, links

def extract_text_and_links_from_image(image_path):
    """Extracts text from image files using OCR."""
    try:
        img = Image.open(image_path)
        text = pytesseract.image_to_string(img)
        links = re.findall(r'https?://\S+', text)
        return text, links
    except Exception as e:
        raise RuntimeError(f"Error extracting text from image: {e}")

def extract_text_and_links_from_pdf(pdf_path):
    text = ""
    links = []
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page_number, page in enumerate(pdf.pages, start=1):
                print(f"Processing page {page_number}...")
                # Extract text
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"

                # Extract links from annotations
                if page.annots:
                    if isinstance(page.annots, list):
                        for annot in page.annots:
                            uri = annot.get("uri", None)
                            if uri:
                                links.append(uri)
                    elif isinstance(page.annots, dict):
                        for annot_key, annot_value in page.annots.items():
                            uri = annot_value.get("uri", None)
                            if uri:
                                links.append(uri)
    except Exception as e:
        raise RuntimeError(f"Error during PDF extraction: {e}")
    return text, links

def extract_text_and_links_from_docx(docx_path):
    import re
    from docx import Document
    import docx2txt

    # Primary extraction using docx2txt
    try:
        text_primary = docx2txt.process(docx_path)
    except Exception as e:
        print(f"docx2txt extraction error: {e}")
        text_primary = ""

    # Fallback extraction using python-docx
    text_secondary = ""
    links = []
    try:
        doc = Document(docx_path)
        # Extract text from paragraphs
        for para in doc.paragraphs:
            text_secondary += para.text + "\n"
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_secondary += cell.text + "\n"
        # Extract text from headers and footers
        for section in doc.sections:
            header = section.header
            for para in header.paragraphs:
                text_secondary += para.text + "\n"
            footer = section.footer
            for para in footer.paragraphs:
                text_secondary += para.text + "\n"
        # Extract hyperlinks from relationships
        for rel in doc.part.rels.values():
            target = rel.target_ref
            if target and "http" in target:
                links.append(target)
    except Exception as e:
        print(f"python-docx extraction error: {e}")

    # Combine text, giving priority to the docx2txt extraction.
    combined_text = (text_primary + "\n" + text_secondary).strip()

    # Remove duplicate lines to avoid overlap
    deduped_lines = []
    for line in combined_text.splitlines():
        line = line.strip()
        if line and line not in deduped_lines:
            deduped_lines.append(line)
    combined_text = "\n".join(deduped_lines)

    # Also extract any additional links from the secondary text
    extra_links = re.findall(r'https?://\S+', text_secondary)
    links.extend(extra_links)
    links = list(set(links))  # Remove duplicates

    return combined_text, links

def extract_text_and_links_from_txt(txt_path):
    try:
        with open(txt_path, "r", encoding="utf-8") as file:
            text = file.read()
            links = re.findall(r'(https?://\S+)', text)
            return text, links
    except Exception as e:
        raise RuntimeError(f"Error extracting text from TXT: {e}")
    
def clean_text(text):
    """
    Cleans the extracted text by removing bullet points and other formatting artifacts.
    """
    # Remove common bullet point characters at the start of lines
    cleaned_lines = []
    for line in text.splitlines():
        # Remove bullet points like '-', '*', '•', '▪', etc.
        cleaned_line = re.sub(r'^(\s*[-*•▪]\s+)', '', line)
        # Optionally, remove numbering (e.g., '1. ', '2) ', etc.)
        cleaned_line = re.sub(r'^(\s*\d+[\.\)]\s+)', '', cleaned_line)
        cleaned_lines.append(cleaned_line)
    
    # Join the cleaned lines back into a single string
    cleaned_text = '\n'.join(cleaned_lines)
    
    # # Optionally, further clean up whitespace
    cleaned_text = re.sub(r'\s+', ' ', cleaned_text).strip()
    
    return cleaned_text

def save_text_and_links_to_file(text, links, output_file_path):
    try:
        with open(output_file_path, "w", encoding="utf-8") as output_file:
            output_file.write("Extracted Text:\n")
            output_file.write(text + "\n\n")
            output_file.write("Extracted Links:\n")
            output_file.write("\n".join(links) + "\n")
        print(f"Extracted text and links saved to: {output_file_path}")
    except Exception as e:
        raise RuntimeError(f"Error saving text and links to file: {e}")

if __name__ == "__main__":
    file_path = input("Enter the file path for the file to extract text and links from: ").strip()

    if not os.path.exists(file_path):
        print(f"File not found: {file_path}")
    else:
        try:
            print("Extracting text and links...")
            extracted_text, extracted_links = extract_text_and_links_from_file(file_path)

            print("Extracted Text Preview:")
            print(extracted_text[:500])
            print("\nExtracted Links:")
            print("\n".join(extracted_links))

            output_file_path = os.path.splitext(file_path)[0] + "_extracted_with_links.txt"
            save_text_and_links_to_file(extracted_text, extracted_links, output_file_path)
        except Exception as e:
            print(f"An error occurred: {e}")
