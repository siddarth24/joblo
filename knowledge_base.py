# knowledge_base.py
import random
from typing import List
from langchain.embeddings.base import Embeddings
from langchain.docstore.document import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
import PyPDF2
import pytesseract
from docx import Document as DocxDocument
from pdf2image import convert_from_path
from sklearn.metrics.pairwise import cosine_similarity
import numpy as np
from langchain.docstore.document import Document as LangDocument
import docx2txt


###########################################################################
# 1) PLACEHOLDER: OllamaEmbeddings
###########################################################################
class OllamaEmbeddings(Embeddings):
    """
    Placeholder for Ollama-based embeddings.
    Replace `_call_ollama_embedding` with your actual logic to call Ollama.
    """

    def __init__(self, model_name: str = "llama2"):
        self.model_name = model_name

    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        return [self._call_ollama_embedding(t) for t in texts]

    def embed_query(self, text: str) -> List[float]:
        return self._call_ollama_embedding(text)

    def _call_ollama_embedding(self, text: str) -> List[float]:
        """
        Implement a real call to Ollama (local server or CLI).
        Currently returns random vectors for demonstration.
        """
        return [random.random() for _ in range(768)]


###########################################################################
# 2) File Parsing Helpers
###########################################################################
def _extract_text_from_pdf(file_path: str) -> str:
    """Extract text from a PDF. Uses OCR for scanned/image-based PDFs."""
    text = ""
    with open(file_path, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"

    # If no text was extracted, use OCR as a fallback
    if not text.strip():
        images = convert_from_path(file_path)  # Convert PDF pages to images
        for img in images:
            ocr_text = pytesseract.image_to_string(img)  # Extract text using OCR
            text += ocr_text + "\n"

    return text.strip()  # Return extracted text


def _extract_text_from_txt(file_path: str) -> str:
    """Extract text from a .txt file."""
    with open(file_path, "r", encoding="utf-8") as f:
        return f.read()


def _extract_text_from_docx(file_path: str) -> str:
    """
    Extract text from a Word document (.docx) using python-docx and a fallback
    with docx2txt to capture additional content (e.g. text in text boxes).
    """
    text_main = ""
    # Primary extraction using python-docx
    try:
        doc = DocxDocument(file_path)
        # Extract text from main paragraphs
        for para in doc.paragraphs:
            text_main += para.text + "\n"
        # Optionally, extract text from tables if present
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_main += cell.text + "\n"
    except Exception as e:
        print(f"python-docx extraction error in {file_path}: {e}")

    # Fallback extraction using docx2txt
    try:
        text_alt = docx2txt.process(file_path)
    except Exception as e:
        print(f"docx2txt extraction error in {file_path}: {e}")
        text_alt = ""

    # Combine the outputs and return the consolidated text
    combined_text = (text_main + "\n" + text_alt).strip()
    return combined_text


###########################################################################
# 3) Build & Retrieve from an In-Memory Vector Store
###########################################################################
def _build_in_memory_vector_store(file_paths: List[str]) -> FAISS:
    """
    Reads each file (PDF, DOCX, or TXT), extracts text (including OCR for PDFs and images in DOCX),
    splits into chunks, embeds with Ollama, and returns an in-memory FAISS vector store.
    """
    docs = []

    # Parse each file
    for path in file_paths:
        if path.lower().endswith(".pdf"):
            text = _extract_text_from_pdf(path)  # Extract text from PDF (including OCR)
        elif path.lower().endswith(".docx"):
            text = _extract_text_from_docx(path)  # Extract text from Word document
        elif path.lower().endswith(".txt"):
            text = _extract_text_from_txt(path)  # Extract text from TXT file
        else:
            continue  # Skip unknown file types

        if text.strip():
            doc = LangDocument(
                page_content=text if text.strip() else "No content extracted",
                metadata={"source": path},
            )
            docs.append(doc)
    # Split into chunks
    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
    chunked_docs = []
    for doc in docs:
        chunks = splitter.split_text(doc.page_content)
        for chunk in chunks:
            chunked_docs.append(Document(page_content=chunk, metadata=doc.metadata))

    # Create embeddings with Ollama
    embeddings = OllamaEmbeddings()

    # Build the FAISS vector store in-memory
    vector_store = FAISS.from_documents(chunked_docs, embeddings)
    return vector_store


def _retrieve_relevant_chunks(vector_store: FAISS, query: str, top_k: int) -> List[str]:
    """
    Given a vector store and a query, retrieve the top_k relevant chunks.
    Returns the chunk contents as a list of strings.
    """
    results = vector_store.similarity_search(query, k=top_k)
    return [doc.page_content for doc in results]


###########################################################################
# 4) Flexible Filtering Function
###########################################################################
RELEVANT_KEYWORDS = [
    "skills",
    "responsibilities",
    "requirements",
    "nice",
    "experience",
    "qualifications",
]


def filter_relevant_keys(data, relevant_keywords=RELEVANT_KEYWORDS):
    """
    Recursively filter a dict (data) to retain only the keys
    whose name or subkeys' names contain any of the keywords
    in `relevant_keywords`.
    """
    if not isinstance(data, dict):
        # If data isn't a dictionary, just return it as is
        return data

    filtered_data = {}

    for key, value in data.items():
        norm_key = key.strip().lower()

        # Check if key itself contains a relevant keyword
        if any(keyword in norm_key for keyword in relevant_keywords):
            # Keep this key. If the value is also a dict, filter it recursively.
            if isinstance(value, dict):
                filtered_value = filter_relevant_keys(value, relevant_keywords)
                filtered_data[key] = filtered_value
            else:
                filtered_data[key] = value

            print(f"Picked relevant key: {key}")

        else:
            # If the key itself isn't relevant, but the value is a dict,
            # we look deeper to see if any nested subkeys might match
            if isinstance(value, dict):
                nested_filtered = filter_relevant_keys(value, relevant_keywords)
                # Only keep the nested dict if it has relevant content
                if nested_filtered:
                    filtered_data[key] = nested_filtered
            else:
                print(f"Ignoring irrelevant key: {key}")

    return filtered_data


###########################################################################
# 5) Re-ranking the retrieved chunks
###########################################################################
def re_rank_chunks(
    chunks: List[str],
    job_description: str,
    embeddings: OllamaEmbeddings,
    top_k: int = 5,
) -> List[str]:
    """
    Given an initial list of retrieved chunks and the job description text,
    re-rank the chunks by their similarity to the job description.

    1. Embed each chunk with OllamaEmbeddings.
    2. Embed the job_description text with OllamaEmbeddings.
    3. Compute cosine similarity between each chunk embedding and the job_description embedding.
    4. Sort by highest similarity and return the top_k chunks.
    """
    if not chunks:
        return []

    # Embed the chunks
    chunk_embeddings = embeddings.embed_documents(chunks)

    # Embed the job description
    job_desc_embedding = embeddings.embed_query(job_description)

    # Compute cosine similarity: shape -> (#chunks, 1)
    sims = cosine_similarity(
        np.array(chunk_embeddings), np.array([job_desc_embedding])
    ).flatten()

    # Sort chunks by similarity score in descending order
    chunk_scores = sorted(zip(chunks, sims), key=lambda x: x[1], reverse=True)

    # Return top_k chunk contents
    top_chunks = [c[0] for c in chunk_scores[:top_k]]
    return top_chunks


###########################################################################
# 6) Main RAG Function to Call from Joblo.py (with Re-Ranking)
###########################################################################
def extract_relevant_chunks(
    file_paths: List[str], job_data: dict, top_k: int = 5, re_rank: bool = True
) -> List[str]:
    """
    1. Filters the job_data to only include relevant keys.
    2. Builds an in-memory FAISS index from PDF/TXT in file_paths.
    3. Dynamically forms a query from the provided filtered job_data dictionary.
    4. Retrieves the top_k relevant chunks.
    5. (Optional) Re-ranks those top_k chunks against the entire job description.
    6. Returns the final chunks as a list of strings.
    """
    if not file_paths:
        return []

    # Filter the job_data for relevant keys
    filtered_job_data = filter_relevant_keys(job_data)

    # Build the vector store
    vector_store = _build_in_memory_vector_store(file_paths)

    # Dynamically form a query from filtered_job_data
    query_lines = []
    for key, value in filtered_job_data.items():
        if isinstance(value, list):
            query_lines.append(
                f"{key.capitalize()}: {', '.join(str(v) for v in value)}"
            )
        else:
            query_lines.append(f"{key.capitalize()}: {value}")
    query = "\n".join(query_lines)

    # Retrieve relevant chunks
    retrieved_chunks = _retrieve_relevant_chunks(vector_store, query, top_k=top_k)

    if re_rank and retrieved_chunks:
        # Combine all filtered job data into one text for re-ranking
        # e.g. "skills: python, c++, requirements: 3 years experience" ...
        job_description_text = " ".join(
            f"{k}: {v if not isinstance(v, list) else ', '.join(str(x) for x in v)}"
            for k, v in filtered_job_data.items()
        )

        embeddings = OllamaEmbeddings()
        # Now re-rank these chunks against the full job description text
        final_chunks = re_rank_chunks(
            retrieved_chunks, job_description_text, embeddings, top_k=top_k
        )
        return final_chunks
    else:
        # If no re-ranking needed or empty retrieval, just return the original top_k chunks
        return retrieved_chunks
